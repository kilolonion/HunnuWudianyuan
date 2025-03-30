import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from pathlib import Path
import re
import os
import tempfile
import base64
import io
import json
import datetime
import hashlib
import importlib.metadata
import shutil

# 检查OpenAI版本
try:
    openai_version = importlib.metadata.version("openai")
    is_old_api = openai_version.startswith("0.")
except:
    # 如果无法确定版本，假定为新版API
    is_old_api = False

import openai
from concurrent.futures import ThreadPoolExecutor

# 默认配置
DEFAULT_CONFIG = {
    "title_keywords": ["举办", "开展", "协助", "组织", "召开", "举行", "宣讲会", "志愿活动", "培训会", "竞赛"],
    "image_keywords": [
        "主持人", "发言", "授课", "讲解", "接听电话", "评分", "作答", "展示", "分享经验",
        "认真听讲", "工作人员合影", "志愿者", "选手演讲", "选手展示", "主讲人", "合影"
    ],
    "redundant_keywords": ["发布人", "浏览数", "日期"],
    "ai_settings": {
        "api_key": "",
        "model": "gpt-3.5-turbo",
        "api_base": ""
    },
    "formatting": {
        "font_name": "宋体",
        "font_size": 12,
        "indent": True
    }
}

# 获取配置目录
def get_config_dir():
    if 'config_dir' in st.session_state and st.session_state.config_dir:
        config_dir = st.session_state.config_dir
    else:
        # 默认配置目录
        config_dir = os.path.join(os.environ.get('APPDATA', 'C:\\'), 'WordFormatter')
    
    # 确保目录存在
    if not os.path.exists(config_dir):
        try:
            os.makedirs(config_dir)
        except Exception as e:
            st.error(f"无法创建配置目录: {str(e)}")
            return None
    
    return config_dir

# 获取配置文件路径
def get_config_path():
    config_dir = get_config_dir()
    if not config_dir:
        return None
    return os.path.join(config_dir, 'config.json')

# 保存配置
def save_config(config):
    config_path = get_config_path()
    if not config_path:
        return False
    
    try:
        with open(config_path, 'w', encoding='utf-8') as f:
            json.dump(config, f, ensure_ascii=False, indent=2)
        return True
    except Exception as e:
        st.error(f"保存配置失败: {str(e)}")
        return False

# 加载配置
def load_config():
    config_path = get_config_path()
    if not config_path or not os.path.exists(config_path):
        # 如果配置文件不存在，创建默认配置
        default_config = DEFAULT_CONFIG.copy()
        save_config(default_config)
        return default_config
    
    try:
        with open(config_path, 'r', encoding='utf-8') as f:
            config = json.load(f)
        
        # 检查并填充缺失的配置项
        for key, value in DEFAULT_CONFIG.items():
            if key not in config:
                config[key] = value
            elif isinstance(value, dict):
                for sub_key, sub_value in value.items():
                    if sub_key not in config[key]:
                        config[key][sub_key] = sub_value
        
        return config
    except Exception as e:
        st.error(f"加载配置失败: {str(e)}")
        return DEFAULT_CONFIG.copy()

# 初始化会话状态
def init_session_state():
    # 加载配置
    config = load_config()
    
    # 设置会话状态
    if 'title_keywords' not in st.session_state:
        st.session_state.title_keywords = config['title_keywords']
    if 'image_keywords' not in st.session_state:
        st.session_state.image_keywords = config['image_keywords']
    if 'redundant_keywords' not in st.session_state:
        st.session_state.redundant_keywords = config['redundant_keywords']
    if 'enable_ai' not in st.session_state:
        st.session_state.enable_ai = False
    if 'api_key' not in st.session_state:
        st.session_state.api_key = config['ai_settings']['api_key']
    if 'model' not in st.session_state:
        st.session_state.model = config['ai_settings']['model']
    if 'api_base' not in st.session_state:
        st.session_state.api_base = config['ai_settings']['api_base']
    if 'font_name' not in st.session_state:
        st.session_state.font_name = config['formatting']['font_name']
    if 'font_size' not in st.session_state:
        st.session_state.font_size = config['formatting']['font_size']
    if 'indent' not in st.session_state:
        st.session_state.indent = config['formatting']['indent']

# 更新配置
def update_config():
    config = {
        "title_keywords": st.session_state.title_keywords,
        "image_keywords": st.session_state.image_keywords,
        "redundant_keywords": st.session_state.redundant_keywords,
        "ai_settings": {
            "api_key": st.session_state.api_key,
            "model": st.session_state.model,
            "api_base": st.session_state.api_base
        },
        "formatting": {
            "font_name": st.session_state.font_name,
            "font_size": st.session_state.font_size,
            "indent": st.session_state.indent
        }
    }
    return save_config(config)

# 标题关键词
TITLE_KEYWORDS = ["举办", "开展", "协助", "组织", "召开", "举行", "宣讲会", "志愿活动", "培训会", "竞赛"]

# 图片说明关键词（用于剔除）
IMAGE_CAPTION_KEYWORDS = [
    "主持人", "发言", "授课", "讲解", "接听电话", "评分", "作答", "展示", "分享经验",
    "认真听讲", "工作人员合影", "志愿者", "选手演讲", "选手展示", "主讲人", "合影"
]

# 系统冗余关键词
REDUNDANT_KEYWORDS = ["发布人", "浏览数", "日期"]

def is_image_caption(text, image_keywords=None):
    if image_keywords is None:
        image_keywords = st.session_state.image_keywords if hasattr(st.session_state, 'image_keywords') else DEFAULT_CONFIG["image_keywords"]
    return any(k in text for k in image_keywords) and len(text) <= 20

def is_redundant(text):
    redundant_kw = st.session_state.redundant_keywords if hasattr(st.session_state, 'redundant_keywords') else DEFAULT_CONFIG["redundant_keywords"]
    return any(k in text for k in redundant_kw)

def is_title(text, title_keywords=None):
    if title_keywords is None:
        title_keywords = st.session_state.title_keywords if hasattr(st.session_state, 'title_keywords') else DEFAULT_CONFIG["title_keywords"]
    return any(k in text for k in title_keywords) and len(text) <= 40

def normalize_review_info(text):
    pattern = re.compile(r"(一审|二审|三审)[：: ]?\s*([\u4e00-\u9fa5]{2,})")
    return [f"{label}：{name}" for label, name in pattern.findall(text)]

def set_style(p, font_name="宋体", font_size=12, bold=False, indent=True, align_left=False):
    run = p.runs[0] if p.runs else p.add_run()
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    if run._element.rPr is not None:
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    p.paragraph_format.first_line_indent = Pt(0 if not indent else 21)
    p.paragraph_format.line_spacing = 1.5
    if align_left:
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

def process_docx(input_path, output_path, title_keywords=None, image_keywords=None,
                font_name="宋体", font_size=12, indent=True, progress_callback=None):
    doc = Document(input_path)
    new_doc = Document()
    seen_titles = set()
    
    # 使用传入的关键词或默认值
    title_kw = title_keywords if title_keywords else TITLE_KEYWORDS
    image_kw = image_keywords if image_keywords else IMAGE_CAPTION_KEYWORDS
    
    total_paragraphs = len(doc.paragraphs)

    for i, para in enumerate(doc.paragraphs):
        # 更新进度
        if progress_callback and total_paragraphs > 0:
            progress_value = int((i / total_paragraphs) * 100)
            progress_callback(progress_value, f"处理段落 {i+1}/{total_paragraphs}")
        
        text = para.text.strip()
        if not text:
            continue
            
        # 检查是否为图片说明
        if is_image_caption(text, image_kw):
            continue
            
        if is_redundant(text):
            continue
            
        if text.startswith("[物电院]"):
            if text not in seen_titles:
                seen_titles.add(text)
                p = new_doc.add_paragraph(text)
                set_style(p, font_name="黑体", font_size=16, bold=True, indent=False, align_left=True)
            continue
            
        # 检查是否为标题
        if is_title(text, title_kw):
            tagged = f"[物电院] {text}"
            if tagged not in seen_titles:
                seen_titles.add(tagged)
                p = new_doc.add_paragraph(tagged)
                set_style(p, font_name="黑体", font_size=16, bold=True, indent=False, align_left=True)
            continue
            
        if text.startswith("（通讯员"):
            p = new_doc.add_paragraph(text)
            set_style(p, indent=False)
            continue
            
        if any(k in text for k in ["一审", "二审", "三审"]):
            for line in normalize_review_info(text):
                p = new_doc.add_paragraph(line)
                set_style(p, indent=False)
            continue
            
        # 普通正文
        p = new_doc.add_paragraph(text)
        set_style(p, font_name=font_name, font_size=font_size, indent=indent)
    
    # 保存文件
    if progress_callback:
        progress_callback(95, "正在保存文件...")
    
    new_doc.save(output_path)

    if progress_callback:
        progress_callback(100, "处理完成")

def extract_docx_text(docx_file):
    """
    从docx文件中提取文本内容用于预览
    """
    if isinstance(docx_file, str):  # 如果是文件路径
        doc = Document(docx_file)
    else:  # 如果是上传的文件对象
        doc = Document(io.BytesIO(docx_file.getvalue()))
    
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    
    return paragraphs

def render_preview(paragraphs, max_height=400):
    """
    渲染文档预览
    """
    if not paragraphs:
        st.info("无内容可预览")
        return
    
    # 创建一个固定高度的容器，带滚动条
    preview_container = st.container()
    
    # 在容器中使用自定义CSS创建一个可滚动的区域
    scrollable_text = f"""
    <div style="height: {max_height}px; overflow-y: auto; border: 1px solid #e6e6e6; padding: 15px; border-radius: 5px; background-color: #f9f9f9;">
    """
    
    # 添加段落
    for para in paragraphs[:100]:  # 限制最多显示100段，避免过大
        if para.startswith("[物电院]") or (len(para) <= 40 and any(kw in para for kw in TITLE_KEYWORDS)):
            # 标题样式
            scrollable_text += f'<p style="font-weight: bold; font-size: 16px; margin-bottom: 8px;">{para}</p>'
        else:
            # 普通段落样式
            scrollable_text += f'<p style="margin-bottom: 8px; text-indent: 2em;">{para}</p>'
    
    if len(paragraphs) > 100:
        scrollable_text += '<p style="color: #888;">...</p>'
    
    scrollable_text += "</div>"
    
    # 显示预览
    preview_container.markdown(scrollable_text, unsafe_allow_html=True)

def get_binary_file_downloader_html(bin_file, file_label='文件'):
    with open(bin_file, 'rb') as f:
        data = f.read()
    b64 = base64.b64encode(data).decode()
    href = f'<a href="data:application/octet-stream;base64,{b64}" download="{os.path.basename(bin_file)}">{file_label}</a>'
    return href

def process_single_file(uploaded_file, title_keywords, image_keywords, font_name, font_size, indent):
    if uploaded_file is None:
        return None, None
        
    # 创建临时文件
    temp_input = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    temp_input.write(uploaded_file.getvalue())
    temp_input.close()
    
    # 创建输出文件路径
    temp_output = tempfile.NamedTemporaryFile(suffix='_标准化处理.docx', delete=False)
    temp_output.close()
    
    # 处理函数
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    def update_progress(value, message=""):
        progress_bar.progress(value / 100)
        status_text.text(message)
    
    try:
        process_docx(
            temp_input.name,
            temp_output.name,
            title_keywords=title_keywords,
            image_keywords=image_keywords,
            font_name=font_name,
            font_size=font_size,
            indent=indent,
            progress_callback=update_progress
        )
        
        # 确保输出文件存在
        if not os.path.exists(temp_output.name):
            st.error("处理后的文件未成功生成")
            return None, None
            
        # 返回处理后的文件和输出文件路径
        output_paragraphs = extract_docx_text(temp_output.name)
        
        # 添加调试信息
        if not output_paragraphs:
            st.warning("处理后的文档内容为空，请检查处理逻辑")
            
        return temp_output.name, output_paragraphs
    except Exception as e:
        st.error(f"处理出错: {str(e)}")
        import traceback
        st.error(f"详细错误: {traceback.format_exc()}")
        return None, None
    finally:
        # 清理临时输入文件
        if os.path.exists(temp_input.name):
            os.unlink(temp_input.name)

def process_batch_files(uploaded_files, title_keywords, image_keywords, font_name, font_size, indent):
    if not uploaded_files:
        return None
    
    # 创建临时目录
    temp_dir = tempfile.mkdtemp()
    output_files = []
    
    # 批量处理进度条
    batch_progress = st.progress(0)
    file_progress = st.progress(0)
    status_text = st.empty()
    
    for i, uploaded_file in enumerate(uploaded_files):
        # 更新批量处理进度
        batch_value = int((i / len(uploaded_files)) * 100)
        batch_progress.progress(batch_value / 100)
        status_text.text(f"处理文件 {i+1}/{len(uploaded_files)}: {uploaded_file.name}")
        
        # 创建临时文件
        temp_input = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        temp_input.write(uploaded_file.getvalue())
        temp_input.close()
        
        # 创建输出文件路径
        output_filename = Path(uploaded_file.name).stem + "_标准化处理.docx"
        output_path = os.path.join(temp_dir, output_filename)
        
        # 文件处理进度回调
        def update_file_progress(value, message=""):
            file_progress.progress(value / 100)
        
        try:
            process_docx(
                temp_input.name,
                output_path,
                title_keywords=title_keywords,
                image_keywords=image_keywords,
                font_name=font_name,
                font_size=font_size,
                indent=indent,
                progress_callback=update_file_progress
            )
            output_files.append((output_filename, output_path))
        except Exception as e:
            st.error(f"处理文件 {uploaded_file.name} 失败: {str(e)}")
        finally:
            # 清理临时输入文件
            os.unlink(temp_input.name)
    
    # 完成
    batch_progress.progress(1.0)
    status_text.text("批处理完成")
    
    return output_files

def create_zip_of_files(files):
    import zipfile
    
    # 创建临时zip文件
    temp_zip = tempfile.NamedTemporaryFile(suffix='.zip', delete=False)
    temp_zip.close()
    
    # 添加文件到zip
    with zipfile.ZipFile(temp_zip.name, 'w') as zipf:
        for filename, filepath in files:
            zipf.write(filepath, arcname=filename)
    
    return temp_zip.name

def extract_content_for_ai(docx_file):
    """
    提取文档内容，用于AI分析
    """
    if isinstance(docx_file, str):  # 如果是文件路径
        doc = Document(docx_file)
    else:  # 如果是上传的文件对象
        doc = Document(io.BytesIO(docx_file.getvalue()))
    
    # 提取前3000个字符用于分析
    content = ""
    for para in doc.paragraphs:
        content += para.text + "\n"
        if len(content) > 3000:
            content = content[:3000]
            break
    
    return content

def analyze_with_openai(content, api_key, model, api_base=None):
    """
    使用OpenAI API分析文档内容，提取关键词
    """
    try:
        # 设置API密钥
        openai.api_key = api_key
        
        # 准备客户端参数
        client_params = {"api_key": api_key}
        
        # 如果提供了自定义API基础URL，则设置它
        if api_base and api_base.strip():
            if is_old_api:
                openai.api_base = api_base
            else:
                client_params["base_url"] = api_base
        
        # 增强提示信息
        prompt = f"""
        你是一位专业的文档分析师，擅长分析学术报告和活动文档。你的任务是从以下文档内容中提取**关键词**，这将用于文档格式化和规范化。

        ## 分析要求
        请识别并提取以下三类关键词：

        1. **标题关键词**：这些词通常出现在文档的标题和小标题中，用于描述具体活动或事件。它们一般是**动词+名词**的组合，标识了文档的核心内容。请提取与活动描述直接相关的关键词，避免长句或短语：
           - 例如：举办、开展、协助、组织、召开、宣讲会、志愿活动、竞赛等。

        2. **图片说明关键词**：这些词通常用于描述图片的内容，简短且与图片直接相关。关键词通常为**动词或名词**，而非长短语。请提取与图片动作或场景相关的词汇：
           - 例如：主持人、发言、展示、合影、授课、讲解等。

        3. **系统冗余关键词**：这些是自动生成的无实质意义的词，通常包括元数据或格式标记，应当被移除：
           - 例如：发布人、浏览数、日期、审稿信息（如一审、二审、三审）等。

        ## 文档上下文
        这份文档是一个{guess_document_type(content)}。请根据文档类型调整你的分析策略。

        ## 文档内容开始：
        {content}
        ## 文档内容结束

        ## 输出要求
        1. 每类关键词至少提供**5个**，最多**15个**。
        2. 关键词应当是**具体**且**简短**，避免长句或描述。
        3. 关键词应当是文档中**实际出现过的**或**高度相关**的词汇。
        4. 严格按照以下JSON格式返回结果，确保格式正确：

        ```json
        {{
          "title_keywords": ["关键词1", "关键词2", ...],
          "image_keywords": ["关键词1", "关键词2", ...],
          "redundant_keywords": ["关键词1", "关键词2", ...]
        }}
        ```

        只返回JSON数据，不要有其他任何解释或说明。
        """

        
        # 根据API版本调用不同的方法
        if is_old_api:
            # 旧版API (openai < 1.0.0)
            # 使用旧版API时忽略linter警告
            # noinspection PyUnresolvedReferences
            response = openai.ChatCompletion.create(
                model=model,
                messages=[
                    {"role": "system", "content": "你是一个专业的文档分析助手，擅长提取文档中的关键信息。你的回答应当简洁、准确、实用，且始终返回有效的JSON格式数据。"},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.2,  # 降低随机性，提高精确度
                max_tokens=1500
            )
            result = response.choices[0].message.content
        else:
            # 新版API (openai >= 1.0.0)
            client = openai.OpenAI(**client_params)
            response = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": "你是一个专业的文档分析助手，擅长提取文档中的关键信息。你的回答应当简洁、准确、实用，且始终返回有效的JSON格式数据。"},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.2,  # 降低随机性，提高精确度
                max_tokens=1500
            )
            result = response.choices[0].message.content
        
        try:
            # 直接尝试解析整个响应为JSON
            try:
                keywords = json.loads(result)
                return keywords
            except:
                # 如果整个响应不是JSON，尝试提取JSON部分
                json_start = result.find('{')
                json_end = result.rfind('}') + 1
                if json_start >= 0 and json_end > json_start:
                    json_str = result[json_start:json_end]
                    keywords = json.loads(json_str)
                    return keywords
                else:
                    st.warning("AI返回的结果不包含有效的JSON数据")
                    return None
        except Exception as e:
            st.warning(f"解析AI返回的JSON数据失败: {str(e)}")
            return None
    except Exception as e:
        st.error(f"调用OpenAI API失败: {str(e)}")
        return None

def guess_document_type(content):
    """
    根据内容推测文档类型
    """
    content_lower = content.lower()
    
    if "培训" in content_lower or "讲座" in content_lower or "报告会" in content_lower:
        return "培训或讲座活动报告"
    elif "竞赛" in content_lower or "比赛" in content_lower:
        return "竞赛活动报告"
    elif "志愿" in content_lower or "公益" in content_lower:
        return "志愿服务活动报告"
    elif "会议" in content_lower:
        return "会议纪要"
    elif "通知" in content_lower or "公告" in content_lower:
        return "通知公告"
    else:
        return "学术活动或机构报告"

def main():
    st.set_page_config(
        page_title="Word文档格式规范工具",
        page_icon="📄",
        layout="wide"
    )
    
    # 应用标题
    st.title("Word文档格式规范工具")
    st.markdown("---")
    
    # 初始化会话状态
    init_session_state()

    # 初始化聊天会话状态
    if 'chat_messages' not in st.session_state:
        st.session_state.chat_messages = []
    if 'chat_token_buffer' not in st.session_state:
        st.session_state.chat_token_buffer = ""
    
    # 侧边栏 - 选项设置
    with st.sidebar:
        # 创建选项卡
        sidebar_tab1, sidebar_tab2, sidebar_tab3, sidebar_tab4 = st.tabs(["基本设置", "AI智能", "系统设置", "AI助手"])
        
        # 基本设置选项卡
        with sidebar_tab1:
            st.header("格式选项")
            
            font_name = st.selectbox(
                "字体",
                options=["宋体", "黑体", "微软雅黑", "仿宋", "楷体"],
                index=0 if st.session_state.font_name not in ["宋体", "黑体", "微软雅黑", "仿宋", "楷体"] else ["宋体", "黑体", "微软雅黑", "仿宋", "楷体"].index(st.session_state.font_name)
            )
            st.session_state.font_name = font_name
            
            font_size = st.select_slider(
                "字体大小",
                options=[10, 12, 14, 16, 18],
                value=st.session_state.font_size
            )
            st.session_state.font_size = font_size
            
            indent = st.checkbox("首行缩进", value=st.session_state.indent)
            st.session_state.indent = indent
            
            st.header("关键词设置")
            
            title_keywords_text = st.text_area(
                "标题关键词",
                value=", ".join(st.session_state.title_keywords),
                height=100
            )
            title_keywords = [kw.strip() for kw in title_keywords_text.split(",") if kw.strip()]
            st.session_state.title_keywords = title_keywords
            
            image_keywords_text = st.text_area(
                "图片说明关键词",
                value=", ".join(st.session_state.image_keywords),
                height=100
            )
            image_keywords = [kw.strip() for kw in image_keywords_text.split(",") if kw.strip()]
            st.session_state.image_keywords = image_keywords
            
            redundant_keywords_text = st.text_area(
                "系统冗余关键词",
                value=", ".join(st.session_state.redundant_keywords),
                height=100
            )
            redundant_keywords = [kw.strip() for kw in redundant_keywords_text.split(",") if kw.strip()]
            st.session_state.redundant_keywords = redundant_keywords
            
            if st.button("保存基本设置"):
                if update_config():
                    st.success("设置已保存到配置文件")
                else:
                    st.error("保存设置失败")
        
        # AI智能选项卡
        with sidebar_tab2:
            st.header("AI智能设置")
            
            enable_ai = st.toggle("启用AI智能关键词分析", value=st.session_state.enable_ai)
            st.session_state.enable_ai = enable_ai
            
            if enable_ai:
                with st.form(key="api_settings"):
                    st.subheader("OpenAI API 设置")
                    
                    api_key = st.text_input(
                        "OpenAI API 密钥",
                        type="password",
                        value=st.session_state.api_key,
                        help="输入你的OpenAI API密钥"
                    )
                    
                    model = st.text_input(
                        "模型名称",
                        value=st.session_state.model,
                        help="输入用于分析的AI模型名称，例如：gpt-3.5-turbo、gpt-4等"
                    )
                    
                    api_base = st.text_input(
                        "API Base URL (可选)",
                        value=st.session_state.api_base,
                        help="适用于使用代理或自定义API端点，留空使用OpenAI默认地址"
                    )
                    
                    submit_button = st.form_submit_button(label="保存AI设置")
                    
                    if submit_button:
                        st.session_state.api_key = api_key
                        st.session_state.model = model
                        st.session_state.api_base = api_base
                        if update_config():
                            st.success("AI设置已保存到配置文件!")
                        else:
                            st.error("保存AI设置失败")
                
                if st.button("测试API连接"):
                    if not st.session_state.api_key:
                        st.error("请先设置API密钥!")
                    else:
                        with st.spinner("正在测试API连接..."):
                            try:
                                openai.api_key = st.session_state.api_key
                                
                                # 准备客户端参数
                                client_params = {"api_key": st.session_state.api_key}
                                if st.session_state.api_base.strip():
                                    if is_old_api:
                                        openai.api_base = st.session_state.api_base
                                    else:
                                        client_params["base_url"] = st.session_state.api_base
                                
                                if is_old_api:
                                    # 旧版API - 忽略linter警告
                                    # noinspection PyUnresolvedReferences
                                    response = openai.ChatCompletion.create(
                                        model=st.session_state.model,
                                        messages=[{"role": "user", "content": "Hello, World!"}],
                                        max_tokens=5
                                    )
                                else:
                                    # 新版API
                                    client = openai.OpenAI(**client_params)
                                    response = client.chat.completions.create(
                                        model=st.session_state.model,
                                        messages=[{"role": "user", "content": "Hello, World!"}],
                                        max_tokens=5
                                    )
                                
                                st.success("API连接测试成功!")
                            except Exception as e:
                                st.error(f"API连接测试失败: {str(e)}")
        
        # 系统设置选项卡
        with sidebar_tab3:
            st.header("系统设置")
            
            config_dir = st.text_input(
                "配置文件目录",
                value=st.session_state.get('config_dir', get_config_dir()),
                help="设置配置文件存储目录，留空使用默认目录"
            )
            
            col1, col2 = st.columns(2)
            with col1:
                if st.button("选择目录"):
                    try:
                        import tkinter as tk
                        from tkinter import filedialog
                        
                        root = tk.Tk()
                        root.withdraw()
                        
                        folder_path = filedialog.askdirectory(
                            title="选择配置文件目录",
                            initialdir=config_dir if os.path.exists(config_dir) else os.path.expanduser("~")
                        )
                        
                        if folder_path:
                            st.session_state.config_dir = folder_path
                            config_dir = folder_path
                            st.rerun()
                    except Exception as e:
                        st.error(f"选择目录失败: {str(e)}")
            
            with col2:
                if st.button("打开配置目录"):
                    try:
                        os.startfile(get_config_dir())
                    except Exception as e:
                        st.error(f"无法打开目录: {str(e)}")
            
            if st.button("应用配置目录"):
                if config_dir and config_dir != get_config_dir():
                    try:
                        # 保存当前配置路径
                        old_config_dir = get_config_dir()
                        
                        # 更新会话状态
                        st.session_state.config_dir = config_dir
                        
                        # 确保新目录存在
                        if not os.path.exists(config_dir):
                            os.makedirs(config_dir)
                        
                        # 如果旧配置存在，复制到新目录
                        old_config_path = os.path.join(old_config_dir, 'config.json')
                        new_config_path = os.path.join(config_dir, 'config.json')
                        
                        if os.path.exists(old_config_path) and not os.path.exists(new_config_path):
                            shutil.copy2(old_config_path, new_config_path)
                        
                        st.success(f"配置目录已更改为: {config_dir}")
                        # 重新加载配置
                        init_session_state()
                    except Exception as e:
                        st.error(f"应用配置目录失败: {str(e)}")
            
            st.markdown("---")
            
            if st.button("恢复默认设置"):
                if st.session_state.get('confirm_reset', False):
                    # 重置为默认配置
                    for key, value in DEFAULT_CONFIG.items():
                        if isinstance(value, dict):
                            for sub_key, sub_value in value.items():
                                if f"{key}_{sub_key}" in st.session_state:
                                    st.session_state[f"{key}_{sub_key}"] = sub_value
                                elif sub_key in st.session_state:
                                    st.session_state[sub_key] = sub_value
                        else:
                            if key in st.session_state:
                                st.session_state[key] = value
                    
                    # 恢复默认配置文件
                    save_config(DEFAULT_CONFIG)
                    
                    st.session_state.confirm_reset = False
                    st.success("已恢复默认设置")
                    st.rerun()
                else:
                    st.session_state.confirm_reset = True
                    st.warning("⚠️ 确定要恢复默认设置吗？再次点击\"恢复默认设置\"确认。")
            else:
                # 重置确认状态
                if 'confirm_reset' in st.session_state:
                    st.session_state.confirm_reset = False
                    
        # AI助手选项卡 - 与AI聊天的功能
        with sidebar_tab4:
            st.header("AI助手")
            
            if not st.session_state.api_key:
                st.warning("请先在「AI智能」设置中配置API密钥")
            else:
                # 显示聊天历史
                chat_container = st.container()
                with chat_container:
                    for msg in st.session_state.chat_messages:
                        with st.chat_message(msg["role"]):
                            st.markdown(msg["content"])
                
                # 用户输入
                user_input = st.chat_input("输入你的问题：")
                
                if user_input:
                    # 添加用户消息到聊天历史
                    st.session_state.chat_messages.append({"role": "user", "content": user_input})
                    
                    # 在界面上显示用户消息
                    with st.chat_message("user"):
                        st.markdown(user_input)
                    
                    # 在界面上添加助手消息占位符
                    with st.chat_message("assistant"):
                        message_placeholder = st.empty()
                    
                    try:
                        # 创建消息列表
                        messages = [{"role": msg["role"], "content": msg["content"]} for msg in st.session_state.chat_messages]
                        
                        # 设置API
                        openai.api_key = st.session_state.api_key
                        client_params = {"api_key": st.session_state.api_key}
                        
                        if st.session_state.api_base.strip():
                            if is_old_api:
                                openai.api_base = st.session_state.api_base
                            else:
                                client_params["base_url"] = st.session_state.api_base
                        
                        full_response = ""
                        
                        # 流式响应
                        if is_old_api:
                            # 旧版API - 忽略linter警告
                            # noinspection PyUnresolvedReferences
                            response = openai.ChatCompletion.create(
                                model=st.session_state.model,
                                messages=messages,
                                stream=True
                            )
                            
                            for chunk in response:
                                if chunk.choices[0].get("delta", {}).get("content"):
                                    content = chunk.choices[0]["delta"]["content"]
                                    full_response += content
                                    message_placeholder.markdown(full_response + "▌")
                        else:
                            # 新版API
                            client = openai.OpenAI(**client_params)
                            stream = client.chat.completions.create(
                                model=st.session_state.model,
                                messages=messages,
                                stream=True
                            )
                            
                            for chunk in stream:
                                if chunk.choices[0].delta.content:
                                    content = chunk.choices[0].delta.content
                                    full_response += content
                                    message_placeholder.markdown(full_response + "▌")
                        
                        # 更新最终响应
                        message_placeholder.markdown(full_response)
                        
                        # 检查是否包含JSON数据并提取
                        try:
                            json_start = full_response.find('{')
                            json_end = full_response.rfind('}') + 1
                            
                            if json_start >= 0 and json_end > json_start:
                                json_str = full_response[json_start:json_end]
                                try:
                                    keywords_data = json.loads(json_str)
                                    
                                    # 检查是否包含关键词字段
                                    if any(k in keywords_data for k in ["title_keywords", "image_keywords", "redundant_keywords"]):
                                        st.success("检测到关键词数据！")
                                        
                                        # 创建应用按钮
                                        if st.button("应用这些关键词"):
                                            # 更新关键词
                                            if "title_keywords" in keywords_data and keywords_data["title_keywords"]:
                                                st.session_state.title_keywords = keywords_data["title_keywords"]
                                            
                                            if "image_keywords" in keywords_data and keywords_data["image_keywords"]:
                                                st.session_state.image_keywords = keywords_data["image_keywords"]
                                            
                                            if "redundant_keywords" in keywords_data and keywords_data["redundant_keywords"]:
                                                st.session_state.redundant_keywords = keywords_data["redundant_keywords"]
                                            
                                            # 保存配置
                                            if update_config():
                                                st.success("成功应用关键词！")
                                                st.rerun()
                                            else:
                                                st.error("保存配置失败")
                                        
                                        # 显示关键词预览
                                        with st.expander("预览关键词"):
                                            if "title_keywords" in keywords_data:
                                                st.write("**标题关键词:**")
                                                st.write(", ".join(keywords_data["title_keywords"]))
                                            
                                            if "image_keywords" in keywords_data:
                                                st.write("**图片说明关键词:**")
                                                st.write(", ".join(keywords_data["image_keywords"]))
                                            
                                            if "redundant_keywords" in keywords_data:
                                                st.write("**系统冗余关键词:**")
                                                st.write(", ".join(keywords_data["redundant_keywords"]))
                                except Exception as e:
                                    st.warning(f"解析JSON失败: {str(e)}")
                        except Exception as e:
                            pass  # 如果没有JSON数据，忽略错误
                            
                        # 添加助手响应到聊天历史
                        st.session_state.chat_messages.append({"role": "assistant", "content": full_response})
                        
                    except Exception as e:
                        st.error(f"发生错误: {str(e)}")
                
                # 清空聊天按钮
                if st.button("清空聊天记录"):
                    st.session_state.chat_messages = []
                    st.rerun()

    # 主界面 - 标签页
    tab1, tab2 = st.tabs(["单文件处理", "批量处理"])
    
    # 创建预览区域容器，在任何选项卡之外，作为共享预览区域
    preview_container = st.container()
    
    # 单文件处理标签页
    with tab1:
        st.header("单文件处理")
        
        uploaded_file = st.file_uploader("选择Word文档", type=["docx"], key="single_file")
        
        if uploaded_file is not None:
            st.write(f"已选择: {uploaded_file.name}")
            
            # 预处理 - 提取原始文档内容
            with st.spinner("加载预览..."):
                input_paragraphs = extract_docx_text(uploaded_file)
            
            # AI分析按钮 - 仅在启用AI和上传文件后显示
            if st.session_state.enable_ai and 'api_key' in st.session_state and st.session_state.api_key:
                if st.button("使用AI分析关键词", key="analyze_ai_single"):
                    with st.spinner("AI正在分析文档..."):
                        content = extract_content_for_ai(uploaded_file)
                        keywords = analyze_with_openai(
                            content, 
                            st.session_state.api_key,
                            st.session_state.model,
                            st.session_state.api_base
                        )
                        
                        if keywords:
                            # 更新会话状态中的关键词
                            if 'title_keywords' in keywords and keywords['title_keywords']:
                                st.session_state.title_keywords = keywords['title_keywords']
                                title_keywords = keywords['title_keywords']
                            
                            if 'image_keywords' in keywords and keywords['image_keywords']:
                                st.session_state.image_keywords = keywords['image_keywords']
                                image_keywords = keywords['image_keywords']
                            
                            if 'redundant_keywords' in keywords and keywords['redundant_keywords']:
                                st.session_state.redundant_keywords = keywords['redundant_keywords']
                                redundant_keywords = keywords['redundant_keywords']
                            
                            st.success("AI分析完成，关键词已更新!")
                            # 显示分析结果
                            with st.expander("查看AI分析结果"):
                                st.write("**标题关键词:**")
                                st.write(", ".join(st.session_state.title_keywords))
                                st.write("**图片说明关键词:**")
                                st.write(", ".join(st.session_state.image_keywords))
                                st.write("**系统冗余关键词:**")
                                st.write(", ".join(st.session_state.redundant_keywords))
            
            # 创建处理按钮和结果容器
            process_btn = st.button("开始处理", key="process_single")
            result_container = st.container()

            output_paragraphs = None  # 初始化输出段落变量
            
            # 在处理按钮点击后处理文档
            if process_btn:
                with st.spinner("处理中..."):
                    output_file, output_paragraphs = process_single_file(
                        uploaded_file,
                        title_keywords,
                        image_keywords,
                        st.session_state.font_name,
                        st.session_state.font_size,
                        st.session_state.indent
                    )
                    
                    if output_file and output_paragraphs:
                        with result_container:
                            st.success("处理完成!")
                            st.markdown(
                                get_binary_file_downloader_html(output_file, '点击下载处理后的文件'),
                                unsafe_allow_html=True
                            )
            
            # 更新预览区域（这部分会在上传文件后立即执行，并在处理完成后再次更新）
            with preview_container:
                update_preview_area(input_paragraphs, output_paragraphs)
    
    # 批量处理标签页
    with tab2:
        st.header("批量处理")
        
        uploaded_files = st.file_uploader("选择多个Word文档", type=["docx"], accept_multiple_files=True, key="batch_files")
        
        batch_input_paragraphs = None
        batch_output_paragraphs = None
        
        if uploaded_files:
            st.write(f"已选择 {len(uploaded_files)} 个文件")
            
            file_list = ""
            for file in uploaded_files:
                file_list += f"- {file.name}\n"
            
            st.markdown(file_list)
            
            # 预览选择的文件
            if len(uploaded_files) > 0:
                preview_file = st.selectbox(
                    "选择要预览的文件",
                    options=[file.name for file in uploaded_files],
                    index=0
                )
                
                # 获取选中的文件对象
                selected_file = next((f for f in uploaded_files if f.name == preview_file), None)
                
                if selected_file:
                    with st.spinner("加载预览..."):
                        batch_input_paragraphs = extract_docx_text(selected_file)
            
            output_files = None  # 初始化输出文件列表
            
            if st.button("开始批量处理", key="process_batch"):
                with st.spinner("批量处理中..."):
                    output_files = process_batch_files(
                        uploaded_files,
                        st.session_state.title_keywords,
                        st.session_state.image_keywords,
                        st.session_state.font_name,
                        st.session_state.font_size,
                        st.session_state.indent
                    )
                    
                    if output_files:
                        st.success(f"批处理完成! 共处理 {len(output_files)} 个文件")
                        
                        # 创建ZIP文件并提供下载
                        zip_file = create_zip_of_files(output_files)
                        st.markdown(
                            get_binary_file_downloader_html(zip_file, '点击下载所有处理后的文件 (ZIP)'),
                            unsafe_allow_html=True
                        )
                        
                        # 更新预览以显示处理后的内容
                        if len(output_files) > 0:
                            preview_output_file = st.selectbox(
                                "选择要预览的处理后文件",
                                options=[filename for filename, _ in output_files],
                                index=0
                            )
                            
                            # 获取选中的文件路径
                            selected_output_path = next((filepath for filename, filepath in output_files if filename == preview_output_file), None)
                            
                            if selected_output_path:
                                batch_output_paragraphs = extract_docx_text(selected_output_path)
                        
                        # 清理临时文件
                        for _, filepath in output_files:
                            if os.path.exists(filepath):
                                os.unlink(filepath)
                        if os.path.exists(zip_file):
                            os.unlink(zip_file)
            
            # 更新批处理预览
            if batch_input_paragraphs:
                with preview_container:
                    update_preview_area(batch_input_paragraphs, batch_output_paragraphs)
    
    # 页脚
    st.markdown("---")
    st.caption("Word文档格式规范工具 © 2023")

def update_preview_area(input_paragraphs, output_paragraphs=None):
    """更新统一的预览区域"""
    st.header("文件预览")
    
    if not input_paragraphs:
        st.info("请上传文件以查看预览")
        return
    
    # 如果没有处理后的内容，只显示输入文件
    if output_paragraphs is None:
        st.subheader("原始文档")
        render_preview(input_paragraphs)
    else:
        # 有处理后的内容，显示对比视图
        st.subheader("文档对比")
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("#### 原始文档")
            render_preview(input_paragraphs, max_height=600)
        
        with col2:
            st.markdown("#### 处理后文档")
            render_preview(output_paragraphs, max_height=600)

if __name__ == "__main__":
    main()
